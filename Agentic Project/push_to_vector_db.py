# push_to_vector_db.py

"""
Simple document processing script for Excel and Word files.
Loads documents and stores them in Postgres with vector embeddings.

Author(s): Saket Khopkar
"""

import os
import pandas as pd
from docx import Document
from dotenv import load_dotenv
from sqlalchemy import create_engine, text
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores.pgvector import PGVector

# ----------------------------
# SETUP
# ----------------------------
load_dotenv()

# Database settings from environment variables
DB_USER = os.getenv("DB_USER")
DB_PASS = os.getenv("DB_PASSWORD")
DB_HOST = os.getenv("DB_HOST")
DB_PORT = os.getenv("DB_PORT")
DB_NAME = os.getenv("DB_NAME")
DB_SCHEMA = os.getenv("DB_SCHEMA", "qpot")

# File paths
EXCEL_FILE = os.path.join("docs", "TestCases.xlsx")
WORD_FILE = os.path.join("docs", "LeaveMgmt.docx")

# Connection string for database
DATABASE_URL = (
    f"postgresql+psycopg2://{DB_USER}:{DB_PASS}@{DB_HOST}:{DB_PORT}/{DB_NAME}"
)
engine = create_engine(DATABASE_URL, isolation_level="AUTOCOMMIT")

# ----------------------------
# HELPER FUNCTIONS
# ----------------------------

def test_database():
    """Check if we can connect to database"""
    try:
        with engine.connect() as conn:
            conn.execute(text("SELECT 1"))
        print("✓ Database connection works")
        return True
    except Exception as e:
        print(f"✗ Database connection failed: {e}")
        return False

def setup_database():
    """Create a fresh vector_documents table in qpot schema"""
    try:
        with engine.begin() as conn:
            conn.execute(text(f"SET search_path TO {DB_SCHEMA}, public"))
            
            # Try to create vector extension
            try:
                conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
                print("✓ Vector extension ready")
            except Exception:
                # Check if extension exists
                result = conn.execute(text("SELECT 1 FROM pg_extension WHERE extname = 'vector'"))
                if not result.fetchone():
                    print("✗ Vector extension not installed. Please ask admin to run: CREATE EXTENSION vector;")
                    return False
                else:
                    print("✓ Vector extension already exists")
            
            # Create fresh table for our documents
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS vector_documents (
                    id SERIAL PRIMARY KEY,
                    doc_name VARCHAR(255) NOT NULL,
                    source_file VARCHAR(255) NOT NULL,
                    chunk_id INTEGER NOT NULL,
                    chunk_text TEXT NOT NULL,
                    embedding VECTOR(384),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    file_type VARCHAR(50)
                )
            """))
            
            # Create indexes for better performance
            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_vector_documents_doc_name 
                ON vector_documents(doc_name)
            """))
            
            conn.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_vector_documents_source_file 
                ON vector_documents(source_file)
            """))
            
            print("✓ Created fresh 'vector_documents' table with indexes")
            print("✓ Database setup complete")
            return True
    except Exception as e:
        print(f"✗ Database setup failed: {e}")
        return False

def load_excel_file(file_path):
    """Load Excel file and convert to text chunks"""
    try:
        df = pd.read_excel(file_path)
        chunks = []
        
        for i, row in df.iterrows():
            # Convert each row to readable text
            text_parts = []
            for column in df.columns:
                if pd.notna(row[column]):  # Skip empty cells
                    text_parts.append(f"{column}: {row[column]}")
            
            if text_parts:  # Only add if row has content
                chunk = "\n".join(text_parts)
                chunks.append(chunk)
        
        print(f"✓ Loaded {len(chunks)} chunks from Excel")
        return chunks
    
    except Exception as e:
        print(f"✗ Error loading Excel: {e}")
        return []

def load_word_file(file_path):
    """Load Word file and convert to text chunks"""
    try:
        doc = Document(file_path)
        chunks = []
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Skip empty paragraphs
                chunks.append(text)
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                
                if row_text:  # Only add if row has content
                    chunks.append(" | ".join(row_text))
        
        print(f"✓ Loaded {len(chunks)} chunks from Word")
        return chunks
    
    except Exception as e:
        print(f"✗ Error loading Word: {e}")
        return []

def save_to_database(chunks, source_name, embeddings):
    """Save text chunks to the fresh vector_documents table"""
    if not chunks:
        print(f"No chunks to save for {source_name}")
        return False
    
    try:
        # Determine file type
        file_type = "excel" if "TestCases" in source_name else "word"
        
        print(f"Generating embeddings for {len(chunks)} chunks...")
        
        saved_count = 0
        batch_size = 20
        
        with engine.begin() as conn:
            conn.execute(text(f"SET search_path TO {DB_SCHEMA}, public"))
            
            for i in range(0, len(chunks), batch_size):
                batch_chunks = chunks[i:i + batch_size]
                
                for j, chunk in enumerate(batch_chunks):
                    try:
                        # Generate embedding for this chunk
                        embedding_vector = embeddings.embed_query(chunk)
                        
                        # Insert into vector_documents table
                        conn.execute(text("""
                            INSERT INTO vector_documents (
                                doc_name, 
                                source_file, 
                                chunk_id, 
                                chunk_text, 
                                embedding,
                                file_type
                            ) VALUES (
                                :doc_name, 
                                :source_file, 
                                :chunk_id, 
                                :chunk_text, 
                                :embedding,
                                :file_type
                            )
                        """), {
                            "doc_name": source_name,
                            "source_file": source_name,
                            "chunk_id": i + j + 1,
                            "chunk_text": chunk,
                            "embedding": embedding_vector,
                            "file_type": file_type
                        })
                        
                        saved_count += 1
                        
                    except Exception as chunk_error:
                        print(f"  ✗ Error saving chunk {i + j + 1}: {chunk_error}")
                        continue
                
                print(f"  Processed batch {i//batch_size + 1} ({len(batch_chunks)} chunks)")
        
        print(f"✓ Saved {saved_count}/{len(chunks)} chunks from {source_name} to vector_documents table")
        return saved_count > 0
    
    except Exception as e:
        print(f"✗ Error saving {source_name}: {e}")
        return False

# ----------------------------
# MAIN PROGRAM
# ----------------------------
def main():
    print("Starting document processing...")
    
    # Step 1: Setup embeddings (this converts text to numbers)
    print("\nSetting up AI embeddings...")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    print("✓ Embeddings ready")
    
    # Step 2: Test database connection
    print("\nTesting database...")
    if not test_database():
        print("Cannot connect to database. Check your .env file!")
        return
    
    # Step 3: Setup database tables
    print("\nSetting up database...")
    if not setup_database():
        print("Cannot setup database!")
        return
    
    # Step 4: Process Excel file
    print("\nProcessing Excel file...")
    if os.path.exists(EXCEL_FILE):
        excel_chunks = load_excel_file(EXCEL_FILE)
        save_to_database(excel_chunks, "TestCases", embeddings)
    else:
        print(f"Excel file not found: {EXCEL_FILE}")
    
    # Step 5: Process Word file
    print("\nProcessing Word file...")
    if os.path.exists(WORD_FILE):
        word_chunks = load_word_file(WORD_FILE)
        save_to_database(word_chunks, "LeaveMgmt", embeddings)
    else:
        print(f"Word file not found: {WORD_FILE}")
    
    # Step 6: Verify data was saved
    print("\nVerifying saved data...")
    try:
        with engine.connect() as conn:
            conn.execute(text(f"SET search_path TO {DB_SCHEMA}, public"))
            
            # Count total rows
            result = conn.execute(text("SELECT COUNT(*) FROM vector_documents"))
            total_count = result.fetchone()[0]
            
            # Count by document type
            result = conn.execute(text("""
                SELECT doc_name, COUNT(*) as count 
                FROM vector_documents 
                GROUP BY doc_name
            """))
            doc_counts = result.fetchall()
            
            print(f"✓ Total chunks in database: {total_count}")
            for doc_name, count in doc_counts:
                print(f"  - {doc_name}: {count} chunks")
                
            # Show table location
            print(f"✓ Data stored in: {DB_SCHEMA}.vector_documents")
            
    except Exception as e:
        print(f"✗ Error verifying data: {e}")
    
    print("\n🎉 All done! Your documents are now in the database.")

if __name__ == "__main__":
    main()
